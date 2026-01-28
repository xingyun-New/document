#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
每日学习计划自动邮件发送程序
功能：每天定时将学习任务转换为Word文档并发送到指定邮箱
"""

import os
import smtplib
import schedule
import time
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("正在安装必要的依赖包...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('study_mailer.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# ==================== 邮箱配置 ====================
EMAIL_CONFIG = {
    'sender_email': 'xingyun1982314@126.com',
    'sender_password': 'MKYurXXfTuE9uZ4p',  # 授权码
    'receiver_email': 'xingyun1982314@126.com',
    'smtp_server': 'smtp.126.com',
    'smtp_port': 465  # 使用SSL加密端口
}

# ==================== 路径配置 ====================
BASE_DIR = Path(__file__).parent
DAILY_TASKS_DIR = BASE_DIR / 'DailyTasks'
OUTPUT_DIR = BASE_DIR / 'output_docs'

# 创建输出目录
OUTPUT_DIR.mkdir(exist_ok=True)

# ==================== 学习计划起始日期 ====================
# Day 1 对应 2026年1月29日
STUDY_START_DATE = datetime(2026, 1, 29)


def get_today_task_file():
    """
    根据当前日期获取对应的学习任务文件
    """
    today = datetime.now()
    
    # 计算今天是第几天
    days_diff = (today - STUDY_START_DATE).days + 1
    
    if days_diff < 1 or days_diff > 7:
        logging.warning(f"今天是第{days_diff}天，超出第一周范围（1-7天）")
        return None
    
    # 日期映射
    date_map = {
        1: '1月29日',
        2: '1月30日',
        3: '1月31日',
        4: '2月1日',
        5: '2月2日',
        6: '2月3日',
        7: '2月4日'
    }
    
    date_str = date_map.get(days_diff)
    if not date_str:
        return None
    
    # 查找对应的文件
    task_file = DAILY_TASKS_DIR / f'Day{days_diff}_第{days_diff}天学习任务_{date_str}.md'
    
    if task_file.exists():
        logging.info(f"找到今天的学习任务文件: {task_file.name}")
        return task_file
    else:
        logging.error(f"未找到学习任务文件: {task_file}")
        return None


def markdown_to_word(md_file, output_file):
    """
    将Markdown文件转换为Word文档
    """
    try:
        logging.info(f"开始转换文件: {md_file.name}")
    except Exception:
        logging.info(f"开始转换文件...")
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            continue
        
        # 标题级别1 (# )
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 标题级别2 (## )
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        
        # 标题级别3 (### )
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        
        # 标题级别4 (#### )
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # 无序列表
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            content = line[2:]
            p = doc.add_paragraph(content, style='List Bullet')
        
        # 有序列表
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            content = line[line.index(' ')+1:]
            p = doc.add_paragraph(content, style='List Number')
        
        # 代码块标记（跳过）
        elif line.startswith('```'):
            continue
        
        # 表格行（简单处理，保持原样）
        elif '|' in line:
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        
        # 普通段落
        else:
            p = doc.add_paragraph(line)
    
    # 保存文档
    try:
        # 使用字符串路径而不是Path对象，避免编码问题
        output_path = str(output_file)
        doc.save(output_path)
        logging.info(f"Word文档已保存: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"保存Word文档失败: {str(e)}")
        raise


def send_email_with_attachment(subject, body, attachment_path):
    """
    发送带附件的邮件
    """
    try:
        logging.info("准备发送邮件...")
        
        # 创建邮件对象
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = EMAIL_CONFIG['receiver_email']
        msg['Subject'] = subject
        
        # 添加邮件正文
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        # 添加附件
        if attachment_path and os.path.exists(attachment_path):
            with open(attachment_path, 'rb') as f:
                attachment = MIMEApplication(f.read(), _subtype='docx')
                # 使用中文文件名作为附件名
                today_cn = datetime.now().strftime('%Y年%m月%d日')
                attachment_filename = f'学习任务_{today_cn}.docx'
                attachment.add_header('Content-Disposition', 'attachment', 
                                    filename=('utf-8', '', attachment_filename))
                msg.attach(attachment)
            logging.info(f"附件已添加")
        
        # 连接SMTP服务器并发送邮件
        with smtplib.SMTP_SSL(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port']) as server:
            server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
            server.send_message(msg)
        
        logging.info(f"✅ 邮件发送成功！收件人: {EMAIL_CONFIG['receiver_email']}")
        return True
        
    except Exception as e:
        logging.error(f"❌ 邮件发送失败: {str(e)}")
        return False


def send_today_study_plan():
    """
    发送今天的学习计划
    """
    logging.info("=" * 60)
    logging.info("开始执行每日学习计划发送任务")
    
    try:
        # 获取今天的任务文件
        task_file = get_today_task_file()
        
        if not task_file:
            logging.warning("今天没有对应的学习任务，跳过发送")
            return
        
        # 生成输出文件名（使用英文格式避免编码问题）
        today_str = datetime.now().strftime('%Y年%m月%d日')
        today_str_en = datetime.now().strftime('%Y-%m-%d')  # 英文格式文件名
        output_file = OUTPUT_DIR / f'StudyTask_{today_str_en}.docx'
        
        # 转换为Word文档
        markdown_to_word(task_file, output_file)
        
        # 准备邮件内容
        subject = f'📚 每日学习计划 - {today_str}'
        body = f"""亲爱的同学：

你好！这是你今天的学习计划。

📅 日期：{today_str}
📖 任务文件：{task_file.name}
⏰ 学习时长：8小时

请查看附件中的详细学习任务，按照计划认真完成今天的学习内容。

记住：
✅ 坚持就是胜利！
✅ 每天进步一点点！
✅ 相信自己，你一定可以！

祝学习顺利！💪

---
本邮件由每日学习计划自动发送系统发送
发送时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        # 发送邮件
        success = send_email_with_attachment(subject, body, str(output_file))
        
        if success:
            logging.info("✅ 今日学习计划发送完成！")
        else:
            logging.error("❌ 学习计划发送失败！")
    
    except Exception as e:
        logging.error(f"❌ 执行任务时发生错误: {str(e)}")
    
    logging.info("=" * 60)


def test_send_now():
    """
    立即测试发送（用于测试）
    """
    print("\n" + "=" * 60)
    print("🧪 测试模式：立即发送今天的学习计划")
    print("=" * 60 + "\n")
    
    send_today_study_plan()
    
    print("\n测试完成！请检查您的邮箱。")


def start_scheduler(send_time="08:00"):
    """
    启动定时任务
    """
    logging.info("=" * 60)
    logging.info("📧 每日学习计划自动邮件发送系统启动")
    logging.info(f"📮 收件邮箱: {EMAIL_CONFIG['receiver_email']}")
    logging.info(f"⏰ 每天发送时间: {send_time}")
    logging.info(f"📁 学习任务目录: {DAILY_TASKS_DIR}")
    logging.info("=" * 60)
    
    # 设置定时任务
    schedule.every().day.at(send_time).do(send_today_study_plan)
    
    # 显示下次执行时间
    next_run = schedule.next_run()
    if next_run:
        logging.info(f"⏰ 下次发送时间: {next_run.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 保持程序运行
    try:
        while True:
            schedule.run_pending()
            time.sleep(60)  # 每60秒检查一次
    except KeyboardInterrupt:
        logging.info("\n程序已停止")


def main():
    """
    主函数
    """
    print("\n" + "=" * 60)
    print("📧 每日学习计划自动邮件发送系统")
    print("=" * 60)
    print("\n请选择运行模式：")
    print("1. 立即测试发送（推荐首次使用）")
    print("2. 启动定时发送（每天早上8点发送）")
    print("3. 自定义发送时间")
    print("0. 退出")
    print()
    
    choice = input("请输入选项 (0-3): ").strip()
    
    if choice == '1':
        test_send_now()
    elif choice == '2':
        start_scheduler("08:00")
    elif choice == '3':
        send_time = input("请输入发送时间（格式：HH:MM，如 08:00）: ").strip()
        try:
            # 验证时间格式
            datetime.strptime(send_time, "%H:%M")
            start_scheduler(send_time)
        except ValueError:
            print("❌ 时间格式错误！请使用 HH:MM 格式，如 08:00")
    elif choice == '0':
        print("再见！")
    else:
        print("❌ 无效的选项！")


if __name__ == '__main__':
    main()

