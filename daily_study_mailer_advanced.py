#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
每日学习计划分时段自动邮件发送程序（增强版）
功能：根据时间表，在每个学习时段开始时发送对应的学习内容
"""

import os
import re
import smtplib
import schedule
import time
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("正在安装必要的依赖包...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('study_mailer_advanced.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# ==================== 邮箱配置 ====================
EMAIL_CONFIG = {
    'sender_email': 'xingyun1982314@126.com',
    'sender_password': 'MKYurXXfTuE9uZ4p',
    'receiver_email': 'xingyun1982314@126.com',
    'smtp_server': 'smtp.126.com',
    'smtp_port': 465
}

# ==================== 路径配置 ====================
BASE_DIR = Path(__file__).parent
DAILY_TASKS_DIR = BASE_DIR / 'DailyTasks'
OUTPUT_DIR = BASE_DIR / 'output_docs'
OUTPUT_DIR.mkdir(exist_ok=True)

# ==================== 学习计划起始日期 ====================
STUDY_START_DATE = datetime(2026, 1, 29)

# ==================== 时间段配置 ====================
# 每天的学习时间段（根据Day1的时间表）
TIME_SLOTS = [
    {'time': '08:00', 'title': '英语早读', 'duration': '30分钟'},
    {'time': '08:30', 'title': '数学基础知识', 'duration': '1小时'},
    {'time': '10:00', 'title': '数学刷题', 'duration': '1小时'},
    {'time': '11:00', 'title': '历史学习', 'duration': '1小时'},
    {'time': '14:00', 'title': '英语学习', 'duration': '1小时'},
    {'time': '15:15', 'title': '生物/道德与法治学习', 'duration': '1小时'},
    {'time': '16:15', 'title': '综合练习', 'duration': '1小时'},
    {'time': '19:00', 'title': '错题复习+预习', 'duration': '1小时'},
]


def get_today_task_file():
    """根据当前日期获取对应的学习任务文件"""
    today = datetime.now()
    days_diff = (today - STUDY_START_DATE).days + 1
    
    if days_diff < 1 or days_diff > 7:
        logging.warning(f"今天是第{days_diff}天，超出第一周范围（1-7天）")
        return None
    
    date_map = {
        1: '1月29日', 2: '1月30日', 3: '1月31日', 4: '2月1日',
        5: '2月2日', 6: '2月3日', 7: '2月4日'
    }
    
    date_str = date_map.get(days_diff)
    task_file = DAILY_TASKS_DIR / f'Day{days_diff}_第{days_diff}天学习任务_{date_str}.md'
    
    if task_file.exists():
        try:
            logging.info(f"Found today's task file: {task_file.name}")
        except:
            logging.info("Found today's task file")
        return task_file
    else:
        logging.error("Task file not found")
        return None


def parse_time_slots_from_md(md_file):
    """
    从Markdown文件中解析时间段内容
    返回: [{'time': '08:00', 'title': '英语早读', 'content': [...]}]
    """
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    time_slots = []
    
    # 匹配时间段标题：### ⏰ 08:00-08:30 | 英语早读
    pattern = r'### ⏰ (\d{2}:\d{2})-(\d{2}:\d{2}) \| (.+?)(?=### ⏰|\Z)'
    matches = re.finditer(pattern, content, re.DOTALL)
    
    for match in matches:
        start_time = match.group(1)
        end_time = match.group(2)
        title = match.group(3).strip()
        slot_content = match.group(0)
        
        time_slots.append({
            'time': start_time,
            'end_time': end_time,
            'title': title,
            'content': slot_content
        })
    
    return time_slots


def create_word_for_time_slot(slot_data, output_file):
    """为单个时间段创建Word文档"""
    try:
        doc = Document()
        
        # 标题（避免emoji在某些环境下的问题）
        try:
            heading = doc.add_heading(f"📚 {slot_data['title']}", 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except:
            heading = doc.add_heading(slot_data['title'], 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 时间信息（避免中文格式化问题）
        time_info = doc.add_paragraph()
        time_info.add_run(f"Learning Time: {slot_data['time']} - {slot_data['end_time']}\n").bold = True
        try:
            today_cn = datetime.now().strftime('%Y-%m-%d')
            time_info.add_run(f"Date: {today_cn}\n")
        except:
            time_info.add_run(f"Date: {datetime.now().date()}\n")
        
        doc.add_paragraph("-" * 30)
        
        # 内容
        lines = slot_data['content'].split('\n')
        for line in lines[1:]:  # 跳过标题行
            line = line.rstrip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            try:
                if line.startswith('####'):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line.strip('*')).bold = True
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif '|' in line:
                    p = doc.add_paragraph(line)
                    p.style = 'Normal'
                else:
                    doc.add_paragraph(line)
            except Exception as e:
                # 如果某行处理失败，跳过但记录
                logging.warning(f"Line processing warning: {str(e)[:50]}")
                continue
        
        # 保存（使用安全的路径处理）
        output_path = str(output_file)
        doc.save(output_path)
        logging.info(f"Word document saved successfully")
        return output_path
        
    except Exception as e:
        logging.error(f"Failed to create Word document: {str(e)}")
        raise


def send_email_with_attachment(subject, body, attachment_path=None):
    """发送带附件的邮件"""
    try:
        logging.info("Preparing to send email...")
        
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = EMAIL_CONFIG['receiver_email']
        
        # 安全设置主题（避免编码问题）
        try:
            msg['Subject'] = subject
        except:
            msg['Subject'] = 'Study Task'
        
        # 安全设置正文
        try:
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
        except:
            msg.attach(MIMEText('Please check the attachment.', 'plain', 'utf-8'))
        
        # 安全添加附件
        if attachment_path and os.path.exists(attachment_path):
            try:
                with open(attachment_path, 'rb') as f:
                    attachment = MIMEApplication(f.read(), _subtype='docx')
                    # 使用简单的英文文件名避免编码问题
                    filename = os.path.basename(attachment_path)
                    attachment.add_header('Content-Disposition', 'attachment', 
                                        filename=('utf-8', '', filename))
                    msg.attach(attachment)
                logging.info("Attachment added")
            except Exception as e:
                logging.error(f"Failed to add attachment: {str(e)[:50]}")
                return False
        
        # 发送邮件
        with smtplib.SMTP_SSL(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port']) as server:
            server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
            server.send_message(msg)
        
        logging.info("Email sent successfully!")
        return True
        
    except Exception as e:
        try:
            logging.error(f"Failed to send email: {str(e)[:100]}")
        except:
            logging.error("Failed to send email")
        return False


def send_time_slot_task(slot_index=None):
    """
    发送指定时间段的学习任务
    如果slot_index为None，则根据当前时间自动判断
    """
    logging.info("=" * 60)
    logging.info(f"开始发送学习任务 (时间段: {slot_index})")
    
    try:
        # 获取今天的任务文件
        task_file = get_today_task_file()
        if not task_file:
            logging.warning("今天没有对应的学习任务")
            return
        
        # 解析时间段
        time_slots = parse_time_slots_from_md(task_file)
        
        if not time_slots:
            logging.error("无法解析时间段内容")
            return
        
        # 确定要发送的时间段
        if slot_index is None:
            # 根据当前时间自动判断
            current_time = datetime.now().strftime('%H:%M')
            slot_index = 0
            for i, slot in enumerate(time_slots):
                if slot['time'] <= current_time:
                    slot_index = i
        
        if slot_index >= len(time_slots):
            logging.warning(f"Time slot index out of range: {slot_index}")
            return
        
        slot_data = time_slots[slot_index]
        
        # 生成Word文档
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
        filename = f"Task_{timestamp}_{slot_data['time'].replace(':', '-')}.docx"
        output_file = OUTPUT_DIR / filename
        
        create_word_for_time_slot(slot_data, output_file)
        
        # 准备邮件（使用安全的日期格式）
        try:
            # 尝试使用中文格式
            today_str = datetime.now().strftime('%Y年%m月%d日')
        except:
            # 失败则使用英文格式
            today_str = datetime.now().strftime('%Y-%m-%d')
        
        try:
            subject = f"⏰ {slot_data['title']} - {today_str} {slot_data['time']}"
        except:
            # 如果主题有编码问题，使用简化版本
            subject = f"Study Task - {slot_data['time']}"
        
        try:
            body = f"""你好！

现在是学习时间！

⏰ 当前时段：{slot_data['time']} - {slot_data['end_time']}
📚 学习内容：{slot_data['title']}
📅 日期：{today_str}

请查看附件中的详细学习任务，开始今天的学习吧！

记住：
✅ 保持专注，避免分心
✅ 按照计划认真完成
✅ 及时记录错题和疑问

加油！💪

---
本邮件由学习计划系统自动发送
发送时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        except:
            # 如果正文有编码问题，使用英文版本
            body = f"""Hello!

It's study time!

Time: {slot_data['time']} - {slot_data['end_time']}
Date: {datetime.now().strftime('%Y-%m-%d')}

Please check the attachment for detailed study tasks.

Good luck!

---
Sent by Study Plan System
Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        # 发送邮件
        success = send_email_with_attachment(subject, body, str(output_file))
        
        if success:
            try:
                logging.info(f"Task sent successfully: {slot_data['title']}")
            except:
                logging.info("Task sent successfully")
        else:
            try:
                logging.error(f"Task sending failed: {slot_data['title']}")
            except:
                logging.error("Task sending failed")
    
    except Exception as e:
        logging.error(f"❌ 执行任务时发生错误: {str(e)}")
    
    logging.info("=" * 60)


def schedule_daily_tasks():
    """设置每日定时任务"""
    logging.info("=" * 60)
    logging.info("Time-slot Study Task Email System Started")
    try:
        logging.info(f"Receiver: {EMAIL_CONFIG['receiver_email']}")
    except:
        logging.info("Receiver: configured")
    logging.info("=" * 60)
    
    # 为每个时间段设置定时任务
    for i, slot in enumerate(TIME_SLOTS):
        schedule.every().day.at(slot['time']).do(send_time_slot_task, slot_index=i)
        try:
            logging.info(f"Scheduled: {slot['time']} - {slot['title']}")
        except:
            logging.info(f"Scheduled: {slot['time']}")
    
    logging.info("=" * 60)
    logging.info("All tasks scheduled!")
    
    # 显示下次执行时间
    next_run = schedule.next_run()
    if next_run:
        logging.info(f"⏰ 下次发送时间: {next_run.strftime('%Y-%m-%d %H:%M:%S')}")
    
    logging.info("=" * 60)
    
    # 保持运行
    try:
        while True:
            schedule.run_pending()
            time.sleep(30)  # 每30秒检查一次
    except KeyboardInterrupt:
        logging.info("\n程序已停止")


def test_send_current_slot():
    """立即测试发送当前时间段的任务"""
    print("\n" + "=" * 60)
    print("🧪 测试模式：立即发送当前时间段的学习任务")
    print("=" * 60 + "\n")
    
    send_time_slot_task()
    
    print("\n测试完成！请检查您的邮箱。")


def test_send_all_slots():
    """测试：发送今天所有时间段的任务（调试用）"""
    print("\n" + "=" * 60)
    print("🧪 测试模式：发送今天所有时间段的学习任务")
    print("=" * 60 + "\n")
    
    task_file = get_today_task_file()
    if not task_file:
        print("未找到今天的任务文件")
        return
    
    time_slots = parse_time_slots_from_md(task_file)
    print(f"找到 {len(time_slots)} 个时间段\n")
    
    for i, slot in enumerate(time_slots):
        print(f"{i+1}. {slot['time']} - {slot['title']}")
    
    print("\n" + "=" * 60)
    choice = input("输入要发送的时间段编号（1-{}），或按回车发送所有: ".format(len(time_slots)))
    
    if choice.strip():
        try:
            slot_index = int(choice) - 1
            send_time_slot_task(slot_index)
        except ValueError:
            print("无效的输入")
    else:
        for i in range(len(time_slots)):
            print(f"\n正在发送第 {i+1} 个时间段...")
            send_time_slot_task(i)
            time.sleep(2)  # 避免发送过快


def main():
    """主函数"""
    print("\n" + "=" * 60)
    print("📧 分时段学习任务自动邮件发送系统（增强版）")
    print("=" * 60)
    print("\n请选择运行模式：")
    print("1. 立即发送当前时间段任务（推荐首次测试）")
    print("2. 启动定时发送（按时间表自动发送所有时间段）")
    print("3. 测试发送特定时间段")
    print("4. 查看今天的时间表")
    print("0. 退出")
    print()
    
    choice = input("请输入选项 (0-4): ").strip()
    
    if choice == '1':
        test_send_current_slot()
    elif choice == '2':
        schedule_daily_tasks()
    elif choice == '3':
        test_send_all_slots()
    elif choice == '4':
        task_file = get_today_task_file()
        if task_file:
            time_slots = parse_time_slots_from_md(task_file)
            print("\n今天的学习时间表：")
            print("=" * 60)
            for i, slot in enumerate(time_slots):
                print(f"{i+1}. {slot['time']} - {slot['end_time']} | {slot['title']}")
            print("=" * 60)
        else:
            print("未找到今天的任务文件")
    elif choice == '0':
        print("再见！")
    else:
        print("❌ 无效的选项！")


if __name__ == '__main__':
    main()

